from flask import Flask, render_template, request, jsonify, url_for, send_file, session, redirect
import sqlite3
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
from dotenv import load_dotenv
import os
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from datetime import datetime
import secrets

load_dotenv()

ADMIN_USERNAME = os.environ.get("ADMIN_USERNAME")
ADMIN_PASSWORD_HASH = os.environ.get("ADMIN_PASSWORD_HASH")

app = Flask(__name__)
app.secret_key = "supersecretkey_change_this"

# ================= CONFIG =================

FROM_EMAIL = os.environ.get("FROM_EMAIL")
APP_PASSWORD = os.environ.get("APP_PASSWORD")

PRINCIPAL_EMAIL = "jagadeeswarangmjs@gmail.com"
INCHARGE_EMAIL = "dhaarani0507@gmail.com"

DB = "bookings.db"

# ================= DATABASE =================

def init_db():
    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute("""
CREATE TABLE IF NOT EXISTS bookings (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    hall TEXT,
    date TEXT,
    slot TEXT,
    booking_type TEXT,
    event_name TEXT,
    name TEXT,
    dept TEXT,
    email TEXT,
    purpose TEXT,
    status TEXT,
    rejection_reason TEXT,
    cancellation_reason TEXT,
    cancel_token TEXT
)
    """)
    conn.commit()
    conn.close()

init_db()

# ================= EMAIL =================

def send_email(to, subject, body, attachment=None):
    msg = MIMEMultipart()
    msg["From"] = FROM_EMAIL
    msg["To"] = to
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "html"))

    if attachment and os.path.exists(attachment):
        with open(attachment, "rb") as f:
            part = MIMEApplication(f.read())
            part.add_header(
                "Content-Disposition",
                "attachment",
                filename=os.path.basename(attachment),
            )
            msg.attach(part)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(FROM_EMAIL, APP_PASSWORD)
        server.sendmail(FROM_EMAIL, to, msg.as_string())

@app.route("/cancel-booking/<token>")
def cancel_by_token(token):
    conn = sqlite3.connect(DB)
    c = conn.cursor()

    c.execute("SELECT * FROM bookings WHERE cancel_token=?", (token,))
    row = c.fetchone()

    if not row:
        conn.close()
        return "<h2>Invalid or expired link.</h2>"

    if row[10] != "approved":
        conn.close()
        return "<h2>Booking already cancelled or invalid.</h2>"

    c.execute("UPDATE bookings SET status='revoked' WHERE cancel_token=?", (token,))
    conn.commit()
    conn.close()

    # Optional: notify admin
    send_email(
        PRINCIPAL_EMAIL,
        "Booking Cancelled by User",
        f"<h3>{row[6]} has cancelled their approved booking for {row[1]} on {row[2]}</h3>"
    )

    return "<h2 style='color:red;text-align:center;margin-top:200px;'>Booking Cancelled Successfully</h2>"

@app.route("/approve/<int:id>")
def approve(id):

    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute("SELECT * FROM bookings WHERE id=?", (id,))
    row = c.fetchone()

    if not row or row[10] != "waiting":
        conn.close()
        return "<h2>Invalid or expired request.</h2>"

    c.execute("UPDATE bookings SET status='approved' WHERE id=?", (id,))
    conn.commit()
    conn.close()

    file_name = f"approval_{id}.pdf"

    doc = SimpleDocTemplate(
        file_name,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=150,
        bottomMargin=80
    )

    elements = []
    styles = getSampleStyleSheet()
    normal = styles["Normal"]

    today = datetime.now().strftime("%d-%m-%Y")

    elements.append(Paragraph(f"Date: {today}", normal))
    elements.append(Spacer(1, 0.3 * inch))

    elements.append(Paragraph("<b>Subject: Approval for Seminar Hall Booking</b>", normal))
    elements.append(Spacer(1, 0.4 * inch))

    body_text = f"""
    Dear {row[6]},<br/><br/>
    This is to formally inform you that your request for booking 
    <b>{row[1]}</b> on <b>{row[2]}</b> during <b>{row[3]}</b> 
    for the event titled "<b>{row[5]}</b>" has been approved.<br/><br/>
    You are requested to ensure proper usage of the hall and hand it 
    over in good condition after the event.<br/><br/>
    We wish you all success for your event.
    """

    elements.append(Paragraph(body_text, normal))
    elements.append(Spacer(1, 0.8 * inch))

    signature_path = os.path.join("static", "signature.png")
    seal_path = os.path.join("static", "approved_seal.png")

    signature_img = ""
    seal_img = ""

    if os.path.exists(signature_path):
        signature_img = Image(signature_path, width=2*inch, height=1*inch)

    if os.path.exists(seal_path):
        seal_img = Image(seal_path, width=1.6*inch, height=1.6*inch)

    bottom_table = Table(
        [[signature_img, seal_img]],
        colWidths=[3*inch, 3*inch]
    )

    bottom_table.setStyle(TableStyle([
        ('ALIGN', (0,0), (0,0), 'LEFT'),
        ('ALIGN', (1,0), (1,0), 'RIGHT'),
        ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
    ]))

    elements.append(bottom_table)
    elements.append(Spacer(1, 0.1 * inch))
    elements.append(Paragraph("<b>Principal</b>", normal))
    elements.append(Paragraph("Mangayarkarasi College of Engineering", normal))

    def add_header(canvas_obj, doc_obj):
        width, height = A4
        header_path = os.path.join("static", "header.png")
        if os.path.exists(header_path):
            canvas_obj.drawImage(
                header_path,
                0,
                height - 140,
                width=width,
                height=130,
                preserveAspectRatio=True
            )

    doc.build(elements, onFirstPage=add_header)

    # ---------------- EMAIL WITH CANCEL BUTTON ----------------

    cancel_link = url_for('cancel_by_token', token=row[13], _external=True)

    email_body = f"""
    <h3>Your Booking is Approved</h3>
    <p>Please find the attached official approval letter.</p>
    <br><br>
    <a href="{cancel_link}" 
    style="background:#e74c3c;color:white;padding:10px 20px;border-radius:6px;text-decoration:none;">
    Cancel Booking
    </a>
    """

    send_email(
        row[8],
        "Seminar Hall Booking Approved",
        email_body,
        attachment=file_name
    )

    return "<h2 style='color:green;text-align:center;margin-top:200px;'>Approved & Email Sent</h2>" 
# ================= ROUTES =================

@app.route("/")
def index():
    return render_template("index.html", bg_image="home_bg.jpg")

# ================= ADMIN LOGIN =================

@app.route("/admin-login", methods=["GET", "POST"])
def admin_login():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")

        if username == "hall@mce" and password == "admin123":
            session["admin"] = True
            return redirect("/admin")
        else:
            return render_template("admin_login.html", error="Invalid Credentials")

    return render_template("admin_login.html")

@app.route("/logout")
def logout():
    session.pop("admin", None)
    return redirect("/")

@app.route("/admin")
def admin():
    if not session.get("admin"):
        return redirect("/admin-login")
    return render_template("admin.html")

@app.route("/api/slot-status")
def slot_status():
    hall = request.args.get("hall")
    date = request.args.get("date")

    conn = sqlite3.connect(DB)
    c = conn.cursor()

    c.execute("""
        SELECT slot, status FROM bookings
        WHERE hall=? AND date=?
        AND (status='approved' OR status='waiting')
    """, (hall, date))

    rows = c.fetchall()
    conn.close()

    MORNING_SLOTS = ["9-10 AM", "10-11 AM", "11-12 PM"]
    AFTERNOON_SLOTS = ["12-1 PM", "1-2 PM", "2-3 PM", "3-4 PM"]

    blocked = {}

    for slot, status in rows:

        # Always block the exact slot
        blocked[slot] = status

        # Hourly morning booking
        if slot in MORNING_SLOTS:
            blocked["half_morning"] = status
            blocked["full_day"] = status

        # Hourly afternoon booking
        elif slot in AFTERNOON_SLOTS:
            blocked["half_afternoon"] = status
            blocked["full_day"] = status

        # Half day morning booking
        elif slot == "half_morning":
            blocked["half_morning"] = status
            blocked["full_day"] = status

        # Half day afternoon booking
        elif slot == "half_afternoon":
            blocked["half_afternoon"] = status
            blocked["full_day"] = status

        # Full day booking
        elif slot == "full_day":
            blocked["full_day"] = status
            blocked["half_morning"] = status
            blocked["half_afternoon"] = status

    return jsonify(blocked)
# ================= API BOOKINGS =================

@app.route("/api/bookings")
def get_bookings():
    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute("SELECT * FROM bookings")
    rows = c.fetchall()
    conn.close()
    return jsonify(rows)

# ================= BOOK SLOT =================

@app.route("/api/book", methods=["POST"])
def book():

    cancel_token = secrets.token_urlsafe(32)
    data = request.get_json()

    if not data:
        return jsonify({"message": "Invalid JSON"}), 400

    required = [
        "hall","date","slot","booking_type",
        "event_name","name","dept","email","purpose"
    ]

    for field in required:
        if not data.get(field):
            return jsonify({"message": f"{field} is required"}), 400

    conn = sqlite3.connect(DB)
    c = conn.cursor()

    # ---------------- SLOT CONFLICT LOGIC ---------------- #

    selected_slot = data["slot"]

    MORNING_SLOTS = ["9-10 AM", "10-11 AM", "11-12 PM"]
    AFTERNOON_SLOTS = ["12-1 PM", "1-2 PM", "2-3 PM", "3-4 PM"]

    FULL_DAY = ["Full Day"]
    HALF_MORNING = ["Half Day (Morning)"]
    HALF_AFTERNOON = ["Half Day (Afternoon)"]

    conflicting_slots = []

    if selected_slot in MORNING_SLOTS:
        conflicting_slots = MORNING_SLOTS + FULL_DAY + HALF_MORNING

    elif selected_slot in AFTERNOON_SLOTS:
        conflicting_slots = AFTERNOON_SLOTS + FULL_DAY + HALF_AFTERNOON

    elif selected_slot in HALF_MORNING:
        conflicting_slots = MORNING_SLOTS + FULL_DAY + HALF_MORNING

    elif selected_slot in HALF_AFTERNOON:
        conflicting_slots = AFTERNOON_SLOTS + FULL_DAY + HALF_AFTERNOON

    elif selected_slot in FULL_DAY:
        conflicting_slots = (
            MORNING_SLOTS +
            AFTERNOON_SLOTS +
            HALF_MORNING +
            HALF_AFTERNOON +
            FULL_DAY
        )

    # If somehow slot doesn't match predefined groups
    if not conflicting_slots:
        conflicting_slots = [selected_slot]

    placeholders = ",".join("?" * len(conflicting_slots))

    query = f"""
    SELECT * FROM bookings
    WHERE hall=? AND date=?
    AND slot IN ({placeholders})
    AND (status='approved' OR status='waiting')
    """

    c.execute(query, [data["hall"], data["date"], *conflicting_slots])

    existing = c.fetchone()

    if existing:
        conn.close()
        return jsonify({
            "message": "Slot conflicts with an existing booking (hour/half/full day overlap)"
        }), 400

    # ---------------- INSERT BOOKING ---------------- #

    c.execute("""
    INSERT INTO bookings
    (hall,date,slot,booking_type,event_name,name,dept,email,purpose,status,cancel_token)
    VALUES (?,?,?,?,?,?,?,?,?,?,?)
    """, (
        data["hall"],
        data["date"],
        data["slot"],
        data["booking_type"],
        data["event_name"],
        data["name"],
        data["dept"],
        data["email"],
        data["purpose"],
        "waiting",
        cancel_token
    ))

    booking_id = c.lastrowid
    conn.commit()
    conn.close()

    # ---------------- EMAIL TO ADMIN ---------------- #

    approve_url = url_for('approve', id=booking_id, _external=True)
    reject_url = url_for('reject', id=booking_id, _external=True)

    body = f"""
    <h2>New Seminar Hall Booking Request</h2>
    <p><b>Event:</b> {data['event_name']}</p>
    <p><b>Name:</b> {data['name']}</p>
    <p><b>Department:</b> {data['dept']}</p>
    <p><b>Hall:</b> {data['hall']}</p>
    <p><b>Date:</b> {data['date']}</p>
    <p><b>Session:</b> {data['slot']}</p>
    <p><b>Purpose:</b> {data['purpose']}</p>
    <br><br>
    <a href="{approve_url}" style="background:green;color:white;padding:10px 20px;border-radius:6px;text-decoration:none;">Approve</a>
    &nbsp;
    <a href="{reject_url}" style="background:red;color:white;padding:10px 20px;border-radius:6px;text-decoration:none;">Reject</a>
    """

    send_email(PRINCIPAL_EMAIL, "Booking Request", body)
    send_email(INCHARGE_EMAIL, "Booking Request", body)

    return jsonify({"message": "Request sent for approval"})

# ================= REJECT =================

@app.route("/reject/<int:id>", methods=["GET","POST"])
def reject(id):
    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute("SELECT * FROM bookings WHERE id=?", (id,))
    row = c.fetchone()

    if not row or row[10] != "waiting":
        conn.close()
        return "<h2>Invalid or expired request.</h2>"

    if request.method == "POST":
        reason = request.form.get("reason")

        c.execute("""
        UPDATE bookings
        SET status='rejected', rejection_reason=?
        WHERE id=?
        """, (reason, id))

        conn.commit()
        conn.close()

        send_email(
            row[8],
            "Booking Rejected",
            f"<h3>Your booking was rejected.</h3><p><b>Reason:</b> {reason}</p>"
        )

        return "<h2 style='color:red;text-align:center;margin-top:200px;'>Rejected</h2>"

    conn.close()

    return """
    <form method="POST" style="text-align:center;margin-top:200px;">
    <h3>Enter Rejection Reason</h3>
    <textarea name="reason" required style="width:300px;height:100px;"></textarea><br><br>
    <button type="submit">Submit Rejection</button>
    </form>
    """

# ================= CANCEL =================

@app.route("/cancel/<int:id>", methods=["POST"])
def cancel(id):
    if not session.get("admin"):
        return jsonify({"message": "Unauthorized"}), 403

    conn = sqlite3.connect(DB)
    c = conn.cursor()

    c.execute("""
    UPDATE bookings
    SET status='cancelled'
    WHERE id=?
    """, (id,))

    conn.commit()
    conn.close()

    return jsonify({"message": "Booking cancelled"})

@app.route("/booking/<hall>")
def booking(hall):
    bg_map = {
        "Seminar Hall 1": "hall1.jpeg",
        "Seminar Hall 2": "hall2.jpeg",
        "TECH HIVE": "tech_hive.jpeg"
    }

    return render_template(
        "booking.html",
        hall=hall,
        bg_image=bg_map.get(hall, "default.jpg")
    )

# ================= DOWNLOAD REPORT =================

@app.route("/download-report")
def download_report():
    if not session.get("admin"):
        return redirect("/admin-login")

    conn = sqlite3.connect(DB)
    c = conn.cursor()
    c.execute("SELECT * FROM bookings")
    rows = c.fetchall()
    conn.close()

    document = Document()
    document.add_heading("Seminar Hall Booking Report", level=1)

    table = document.add_table(rows=len(rows)+1, cols=10)
    table.style = "Table Grid"

    headers = [
        "ID","Hall","Date","Slot","Type",
        "Event","Name","Dept","Email","Status"
    ]

    for col, header in enumerate(headers):
        table.rows[0].cells[col].text = header

    for row_index, row in enumerate(rows):
        for col in range(10):
            table.rows[row_index+1].cells[col].text = str(row[col] if col != 9 else row[10])

    file_path = "booking_report.docx"
    document.save(file_path)

    return send_file(file_path, as_attachment=True)

# ================= RUN =================

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)